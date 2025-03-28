import math
import logging
import datetime
import matplotlib.pyplot as plt
from docx import Document
import os
from config import Config

logger = logging.getLogger(__name__)

class LoanProcessor:
    @staticmethod
    def calculate_eligibility_score(work_duration: int, factory_workers: int, 
                                  factory_locations: int, loan_amount: float) -> tuple:
        """Calculate loan eligibility score based on employment factors"""
        if work_duration < 6:
            return -1, "Minimum employment duration is 6 months."
        if factory_workers < 5:
            return -1, "Minimum number of factory workers is 5."
        if factory_locations < 1:
            return -1, "At least 1 factory location is required."
            
        score = (work_duration * 2) + factory_workers + (factory_locations * 0.5) - (loan_amount / 50)
        return score, None

    @staticmethod
    def calculate_loan_terms(loan_amount: float, tenor: str) -> tuple:
        """Calculate loan interest and repayment terms"""
        if tenor.endswith('w'):
            duration_value = int(tenor[:-1])
            total_weeks = duration_value
        elif tenor.endswith('m'):
            duration_value = int(tenor[:-1])
            total_weeks = duration_value * 4
        else:
            duration_value = 1
            total_weeks = 4

        interest_rate = 0.05
        total_interest = loan_amount * interest_rate * duration_value
        processing_charge = 5
        total_repayment = loan_amount + total_interest + processing_charge

        return total_interest, processing_charge, total_repayment, total_weeks

    @staticmethod
    def calculate_payment_schedule(total_repayment: float, total_weeks: int, 
                                 frequency: str) -> tuple:
        """Calculate payment schedule based on frequency"""
        if frequency == "weekly":
            installments = total_weeks
            interval_days = 7
        elif frequency == "biweekly":
            installments = math.ceil(total_weeks / 2)
            interval_days = 14
        else:  # monthly
            installments = math.ceil(total_weeks / 4)
            interval_days = 30

        installment_amount = total_repayment / installments
        schedule = []
        
        for i in range(installments):
            due_date = datetime.date.today() + datetime.timedelta(days=interval_days * (i+1))
            schedule.append({
                'installment': i+1,
                'due_date': due_date,
                'amount': installment_amount
            })

        return installments, installment_amount, schedule

    @staticmethod
    def generate_repayment_chart(installments_list: list, language: str) -> str:
        """Generate visual chart of repayment schedule"""
        installments = [i['installment'] for i in installments_list]
        amounts = [i['amount'] for i in installments_list]
        due_dates = [i['due_date'].strftime('%Y-%m-%d') for i in installments_list]

        plt.figure(figsize=(10, 6))
        plt.bar(installments, amounts, color='skyblue')
        plt.xlabel("Installment Number")
        plt.ylabel("Amount ($)")
        plt.title("Loan Repayment Schedule")
        plt.xticks(installments, due_dates, rotation=45)
        plt.tight_layout()

        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        file_name = f"Repayment_Chart_{timestamp}.png"
        file_path = os.path.join(Config.UPLOAD_DIR, file_name)
        plt.savefig(file_path)
        plt.close()
        
        logger.info(f"Repayment chart generated at {file_path}")
        return file_path

    @staticmethod
    def generate_loan_report(context, user_folder: str) -> str:
        """Generate detailed loan report document"""
        document = Document()
        document.add_heading('Loan Report & Agreement', 0)

        # Contract Details
        document.add_heading('Contract Details', level=1)
        document.add_paragraph(f"Name: {context.user_data.get('user_full_name', 'N/A')}")
        document.add_paragraph(f"Loan Amount: ${context.user_data.get('amount', 'N/A')}")
        document.add_paragraph(f"Loan Tenor: {context.user_data.get('tenor', 'N/A')}")
        document.add_paragraph(f"Processing Charge: ${context.user_data.get('processing_charge', 'N/A')}")
        document.add_paragraph(f"Total Interest: ${context.user_data.get('total_interest', 'N/A')}")

        # Employment Details
        document.add_heading('Employment Details', level=1)
        document.add_paragraph(f"Employment Duration: {context.user_data.get('work_duration', 'N/A')} months")
        document.add_paragraph(f"Factory Workers: {context.user_data.get('factory_workers', 'N/A')}")
        document.add_paragraph(f"Factory Locations: {context.user_data.get('factory_locations', 'N/A')}")
        document.add_paragraph(f"Factory Origin: {context.user_data.get('factory_origin', 'N/A')}")

        # Repayment Details
        document.add_heading('Repayment Details', level=1)
        document.add_paragraph(f"Total Repayment Amount: ${context.user_data.get('total_repayment', 'N/A')}")
        document.add_paragraph(f"Payment Frequency: {context.user_data.get('payment_frequency', 'N/A')}")
        schedule_text = context.user_data.get('schedule_text', 'N/A')
        document.add_paragraph("Repayment Schedule:")
        document.add_paragraph(schedule_text)

        # Save the document
        telegram_id = context.user_data.get('telegram_id', 'report')
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        file_name = f"Loan_Report_{telegram_id}_{timestamp}.docx"
        file_path = os.path.join(user_folder, file_name)
        document.save(file_path)
        
        logger.info(f"Report generated at {file_path}")
        return file_path

    @staticmethod
    def simulate_loan(amount: float, annual_interest_rate: float, term_months: int) -> tuple:
        """Simulate loan payments with given parameters"""
        r = (annual_interest_rate / 100) / 12  # Monthly interest rate
        if r == 0:
            monthly_payment = amount / term_months
        else:
            monthly_payment = amount * (r * (1 + r) ** term_months) / ((1 + r) ** term_months - 1)
        
        total_payment = monthly_payment * term_months
        return monthly_payment, total_payment